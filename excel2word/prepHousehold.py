'''
- occupation: right justified
https://stackoverflow.com/questions/28884114/python-docx-align-both-left-and-right-on-same-line?noredirect=1&lq=1

- GUI:
    - select which elements to include
    - adjust font size

- table style: https://github.com/python-openxml/python-docx/issues/9

-Create a GUI which allows a user to select which fields to include and specify the font size
-Adjust the table style to match sample doc (indentation, cell padding, border color)
-Increase page column width to match sample doc (meh)
    

-create a district-based title page
The title page should include a list of numbered household heads. 
If household numbers are available, use those, and if not, just a 1., 2., 3. is fine. 
The district name should be at the top in larger letters and centered. 
If someone is a deacon, minister, or bishop, those words should follow the household name and in parentheses, and only the last ordination (so bishop trumps minister trumps deacon) 
Also, a header should be at the top and centered with the name of the settlement and the name of the district. 
An option should be made available in the program to exclude the settlement name (in case the whole directory is the same settlement).

-table of contents and index. 
 -we need the church settlement and district title pages to be the main points.
 -then two separate indexes: one for the household head, one for the wife.    
'''

from extractExcel import getData
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER


def formatDate(prefix, hh, z=False):
    d = ''
    if z:
        if hh.contains(prefix + 'm'):
            d += hh.get(prefix + 'm').zfill(2) + '/'
        if hh.contains(prefix + 'd'):
            d += hh.get(prefix + 'd').zfill(2) + '/'
        if hh.contains(prefix + 'y'):
            d += hh.get(prefix + 'y').zfill(2)
    else:
        if hh.contains(prefix + 'm'):
            d += hh.get(prefix + 'm') + '/'
        if hh.contains(prefix + 'd'):
            d += hh.get(prefix + 'd') + '/'
        if hh.contains(prefix + 'y'):
            d += hh.get(prefix + 'y')
    return d
        

### Line 1 - name and occupation
def getLineOne(hh, doc):
    p = doc.add_paragraph('')
    s = ''
    if hh.contains('hhcode'):
        s += hh.get('hhcode') + ' '
    s += hh.get('lastname').upper() + ', ' + hh.get('firstname').upper() + ' ' + hh.get('middle').upper()
    if hh.contains('suffix'):
        s += ', ' + hh.get('suffix').upper()
    if hh.contains('member'):
        if hh.get('member') == 'y':
            s += '*'
    run = p.add_run(removeStraySpaces(s))
    run.bold = True
    f = run.font
    f.size = Pt(11)
    
    if hh.contains('occup1'):
        s = '\t' + hh.get('occup1').title()
        for x in range(2,5):
            if hh.contains('occup' + str(x)):
                s += ' / ' + hh.get('occup' + str(x)).title()
        p.add_run(removeStraySpaces(s))
    
    tab_stops = p.paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES) #...LEADER.TABS also an option
    return doc

### Line 2 - address and contact info
def getLineTwo(hh):
    s = ''
    if hh.contains('address'):
        s += hh.get('address').title() + ', '
    if hh.contains('town'):
        s += hh.get('town').title() + ', ' 
    s += hh.get('state').upper() + ' ' + hh.get('zip') + ' '
    if hh.contains('telephone'):
        s += formatPhone(hh.get('telephone'))
    if hh.contains('email'):
        s += '; ' + hh.get('email')
    return removeStraySpaces(s)
    
### Line 3 - ordinations
def getLineThree(hh):
    s = ''
    if hh.contains('ordain_deac'):
        s += 'Deac. ' + formatDate('ordain_deac', hh) + '; '
    if hh.contains('ordain_mins'):
        s += 'Mins. ' + formatDate('ordain_mins', hh) + '; '
    if hh.contains('ordain_bish'):
        s += 'Bish. ' + formatDate('ordain_bish', hh)
    return removeStraySpaces(s)

### Line 4 - personal info
def getLineFour(hh, doc):
    p = doc.add_paragraph('')
    
    t = ''
    b = ''
    s = 'b. ' + formatDate('born', hh) + ', '
    if hh.contains('hhhdiedm'):
        s += 'd. ' + formatDate('hhhdied', hh) + ', ' 
    if hh.contains('fatherfirst'):
        s += 's.o. ' + hh.get('fatherfirst').title() + ' ' + hh.get('fathermiddle').title() + ' ' + formatSuffix(hh.get('fathersuffix'))
        if hh.contains('motherfirst'):
            s += ' & ' + hh.get('motherfirst').title() + ' ' + hh.get('mothermiddle').title() + ' (' + hh.get('motherlast').title() + ') ' + hh.get('fatherlast').title()
        if hh.contains('hhhparentcode'):
            s += ' [' + hh.get('hhhparentcode') + ']'
    if hh.contains('hhhmovedfrom'):
        s += ', moved in ' + formatDate('hhhmovedfrom', hh) + ' from ' + formatAddress(hh.get('hhhmovedfrom'))
    s += '; '
    if hh.contains('cmary'):
        s += 'm. ' + formatDate('cmar', hh)
    if hh.contains('cwfname'):
        s += ' to ' 
        b = hh.get('cwfname').upper() + ' ' + hh.get('cwmiddle').upper() + ' ' + hh.get('cwlname').upper()
        if hh.get('cwmember') == 'y':
            b += '*'
        if hh.contains('cwborny'):
            t += ', b. ' + formatDate('cwborn', hh)
        if hh.contains('cwdiedy'):
            t += ', d. ' + formatDate('cwdied', hh)
        if hh.contains('cwdadfname'):
            t += ', d.o. ' + hh.get('cwdadfname').title() + ' ' + hh.get('cwdadmname').title() + ' ' + hh.get('cwdadlname').title() + ' ' + formatSuffix(hh.get('cwdadsname'))
            if hh.contains('cwmomfname'):
                t += ' & ' + hh.get('cwmomfname').title() + ' ' + hh.get('cwmommname').title() + ' (' + hh.get('cwmomlname').title() + ') ' + hh.get('cwdadlname').title()
            if hh.contains('cwparentcode'):
                t += ' [' + hh.get('cwparentcode') + ']'
    if hh.contains('cwmovedfrom'):
        t += ', moved in ' + formatDate('cwmovedfrom', hh) + ' from ' + formatAddress(hh.get('cwmovedfrom')) + '.'
    if hh.contains('movedfrom'):
        t += ' Moved in ' + formatDate('movedfrom', hh) + ' from ' + formatAddress(hh.get('movedfrom')) + '.'
    p.add_run(removeStraySpaces(s))
    p.add_run(removeStraySpaces(b)).bold = True
    p.add_run(removeStraySpaces(t))
    return doc
    
### Children
### params: household, document, prefix
def getChildren(hh, doc, p):
    if not hh.contains(p + '01fname'):
        if p == 'cwc':
            doc.add_paragraph('')
        return doc
    table = doc.add_table(rows=1, cols=4)
    table = set_col_widths(table)
    table.style = 'Table Grid'
    
    cells = table.rows[0].cells

    n = 1
    while hh.contains(p + str(n).zfill(2) + 'fname'):
        name = hh.get(p + str(n).zfill(2) + 'fname').title() + ' '
        name += hh.get(p + str(n).zfill(2) + 'mname').title() + ' ' 
        name += formatSuffix(hh.get(p + str(n).zfill(2) + 'sname'))
        cells[0].text = removeStraySpaces(name).strip()
        cells[1].text = formatDate(p + str(n).zfill(2) + 'born', hh, True)
        cells[2].text = hh.get(p + str(n).zfill(2) + 'bc').upper()
        s = ''
        if hh.contains(p + str(n).zfill(2) + 'diedy'):
            s += 'd. ' + formatDate(p + str(n).zfill(2) + 'died', hh, True) + ' '
        s += hh.get(p + str(n).zfill(2) + 'spousefname').title() + ' ' 
        s += hh.get(p + str(n).zfill(2) + 'spousemname').title() + ' ' 
        s += hh.get(p + str(n).zfill(2) + 'spouselname').title() + ' ' 
        s += formatSuffix(hh.get(p + str(n).zfill(2) + 'spousesname')) + ' '
        if len(s.strip()) > 0 and hh.contains(p + str(n).zfill(2) + 'address'):
            s += ': '
        s += formatAddress(hh.get(p + str(n).zfill(2) + 'address')) + ' '
        if hh.contains(p + str(n).zfill(2) + 'hshld#'):
            s += ' [' + hh.get(p + str(n).zfill(2) + 'hshld#') + ']'
        cells[3].text = removeStraySpaces(s).strip()
        n += 1
        if hh.contains(p + str(n).zfill(2) + 'fname'):
            cells = table.add_row().cells
            
    table = set_font_size(table, 8)
    doc.add_paragraph('')
    return doc
        
# Set column widths
def set_col_widths(table):
    widths = (1, .7, .3, 1.65)
    i=0
    for cell in table.columns[0].cells:
        cell.width = Inches(widths[i])
        #TODO: remove padding
        # https://groups.google.com/forum/#!topic/python-docx/ABjObkKkOu0
        i += 1
    i = 0
    for column in table.columns:
        column.width = Inches(widths[i])
        i += 1
    return table

# Set font size for a table
def set_font_size(table, size):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(size)  
    
### Children of other spouses
def getPrevSpouse(hh, doc, p, head=True):
    if not hh.contains(p + 'fname'):
        return doc
    
    par = doc.add_paragraph('')
    s = ''
    b = ''
    t = ''
    if head:
        s += hh.get('firstname').title()
    else:
        s += hh.get('cwfname').title()
    s += ' m. ' + formatDate(p + 'mar', hh) + ' '
    b += hh.get(p + 'fname').upper() + ' ' + hh.get(p + 'mname').upper() + ' ' + hh.get(p + 'lname').upper()
    t += ' b. ' + formatDate(p + 'born', hh)
    if hh.contains(p + 'diedm'):
        t += ' d. ' + formatDate(p + 'died', hh)
    if hh.contains(p + 'dadfname'):
        t += ' d.o. ' + hh.get(p + 'dadfname').title()
        if hh.contains(p + 'momfname'):
            t += ' & ' + hh.get(p + 'momfname').title() + ' (' + hh.get(p + 'momlname').title() + ')'
        t += ' ' + hh.get(p + 'dadlname').title()
        if hh.contains(p + 'parentcode'):
            t += ' [' + hh.get(p + 'parentcode') + ']'
        if hh.contains(p + 'movedfrom'):
            t += ', moved in ' + formatDate(p + 'movedfrom', hh) + ' from ' + formatAddress(hh.get(p + 'movedfrom')) + '.'
    par.add_run(removeStraySpaces(s))
    par.add_run(removeStraySpaces(b)).bold = True
    par.add_run(removeStraySpaces(t))
    doc = getChildren(hh, doc, p)
    return doc
    
def formatPhone(p):
    s = '('
    for c in p:
        if len(s) == 4:
            s += ') '
        if len(s) == 9:
            s += '-'
        if c.isdigit():
            s += c
    return s
    
def formatAddress(s):
    t = ''
    s = s.split()
    for c in s:
       if len(c) == 2 and c.isalpha():
          t += c.upper()
       else:
          t += c.title()
       t += ' '
    return t.strip()
    
def formatSuffix(s):
    if s.lower() == 'jr' or s.lower() == 'sr':
        return s.title()
    else:
        return s.upper()
    
def removeStraySpaces(s):
    while '  ' in s:
        s = s.replace('  ', ' ')
    return s
    