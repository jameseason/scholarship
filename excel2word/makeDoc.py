from extractExcel import getData
from prepHousehold import getLineOne, getLineTwo, getLineThree, getLineFour, getChildren, getPrevSpouse, getHeadName
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
import time
from docx.oxml import OxmlElement
from Tkinter import *
from docx.oxml.ns import qn

w = Tk()
inputLocation = StringVar()
outputLocation = StringVar()
attrRow = StringVar()
startRow = StringVar()
endRow = StringVar()
titleCheck = IntVar()

    
def makeDoc(hhs):
    print "generating document..."
    doc = Document()
    #doc.add_paragraph("first page blank")
    #doc.add_page_break()
    if titleCheck.get() != 1:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        set_number_of_columns(section, 2)
    style = doc.styles['Normal']
    
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    
    current_settlement = None
    current_district = None 
    list = None
    
    count = 1
    
    
    for hh in hhs:
        if titleCheck.get() == 1:
            if hh.get('settlement') == current_settlement and hh.get('chdist') == current_district:
                if hh.contains('hhcode'):
                   list.add_run(hh.get('hhcode') + ' ' + getHeadName(hh) + '\n')
                else:
                   list.add_run(str(count) + '. ' + getHeadName(hh) + '\n')
                count += 1
            else:
                current_settlement = hh.get('settlement')
                current_district = hh.get('chdist')
                if list != None: #if this isn't the first dist
                   doc.add_page_break()
                   sec1 = doc.add_section(WD_SECTION.NEW_PAGE)
                   set_number_of_columns(sec1, 1)
                   
                
                head = doc.add_paragraph(hh.get('settlement').title() + ': ' + hh.get('chdist').title())
                head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                set = doc.add_paragraph('')
                set.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                run = set.add_run(hh.get('chdist').title())
                run.font.size = Pt(20)
                
                count = 1
                if hh.contains('hhcode'):
                   first_hh = hh.get('hhcode') + ' ' + getHeadName(hh) + '\n'
                else:
                   first_hh = str(count) + '. ' + getHeadName(hh) + '\n'
                   
                list = doc.add_paragraph(first_hh)
                count += 1
                
                doc.add_page_break()
                sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
                set_number_of_columns(sec2, 2)            
            
            
            
        doc = getLineOne(hh, doc)
        two = doc.add_paragraph(getLineTwo(hh).strip())
        three = getLineThree(hh).strip()
        if len(three) > 0:
            doc.add_paragraph(three)
        doc = getLineFour(hh, doc)
        doc = getChildren(hh, doc, 'cwc')
        doc = getPrevSpouse(hh, doc, '1w', True)
        doc = getPrevSpouse(hh, doc, '2w', True)
        doc = getPrevSpouse(hh, doc, '1h', False)
        doc = getPrevSpouse(hh, doc, '2h', False)
    
    secs = doc.sections
    for sec in secs:
        sec.top_margin = Cm(.5)
        sec.bottom_margin = Cm(.5)
        sec.left_margin = Cm(.5)
        sec.right_margin = Cm(.5)
    
    try:
        print "saving to " + outputLocation.get()
        doc.save(outputLocation.get())
    except:
        print 'permission denied. trying again in 10'
        time.sleep(10)
        doc.save(outputLocation.get())
        
    
   

def set_number_of_columns(section, cols):
    """ sets number of columns through xpath. """
    WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))
    

if __name__ == "__main__":
    run()
   
def run():
    print 'started'
    hhs = getData(inputLocation.get(), attrRow.get(), startRow.get(), endRow.get())
    print 'ready to make document'
    makeDoc(hhs)
    print 'done'
